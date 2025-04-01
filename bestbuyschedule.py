# Created by Ratheesh Mathew
# For BestBuy Work Schedule  

import calendar
from datetime import datetime
import docx
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def create_month_calendar(year, month):
    # Create a new Word document
    doc = Document()
    
    # Set up the calendar object
    cal = calendar.Calendar(firstweekday=calendar.SUNDAY)
    month_days = cal.monthdayscalendar(year, month)
    
    # Get month name and abbreviations
    month_name = calendar.month_name[month]
    month_abbr = month_name[:3].upper()
    prev_month = month - 1 if month > 1 else 12
    prev_year = year if month > 1 else year - 1
    next_month = month + 1 if month < 12 else 1
    next_year = year if month < 12 else year + 1
    next_month_abbr = calendar.month_name[next_month][:3].upper()
    prev_month_abbr = calendar.month_name[prev_month][:3].upper()
    
    doc.add_heading(f'{month_name} {year}', level=1)
    
    # Create table
    table = doc.add_table(rows=len(month_days) + 1, cols=7)
    table.style = 'Table Grid'
    
    # Determine first day of the month
    first_day = datetime(year, month, 1).weekday()
    first_day = (first_day + 1) % 7  # Adjust to make Sunday=0
    
    # Set column headers
    days = ['SUN', 'MON', 'TUE', 'WED', 'THU', 'FRI', 'SAT']
    header_cells = table.rows[0].cells
    for i, day in enumerate(days):
        if i < first_day:  # Days before month start
            header_cells[i].text = f"{prev_month_abbr}-{day}"
        else:  # Days of current month
            header_cells[i].text = f"{month_abbr}-{day}"
        
        # Apply background color to header row
        shading_elm = parse_xml(r'<w:shd {} w:fill="D3D3D3"/>'.format(nsdecls('w')))
        header_cells[i]._tc.get_or_add_tcPr().append(shading_elm)
    
    # Reduce first row height
    table.rows[0].height = Inches(0.25)
    
    prev_cal = calendar.Calendar(firstweekday=calendar.SUNDAY)
    prev_month_days = prev_cal.monthdayscalendar(prev_year, prev_month)
    
    for week_num, week in enumerate(month_days, 1):
        for day_num, day in enumerate(week):
            cell = table.rows[week_num].cells[day_num]
            paragraph = cell.paragraphs[0]
            
            if day == 0 and week_num == 1:  # Previous month's days
                prev_week = prev_month_days[-1]
                paragraph.add_run(f"{prev_month_abbr}-{prev_week[day_num]}\n")
                shading_elm = parse_xml(r'<w:shd {} w:fill="D3D3D3"/>'.format(nsdecls('w')))
                cell._tc.get_or_add_tcPr().append(shading_elm)
            elif day == 0 and week_num == len(month_days):  # Next month's days
                next_day = day_num + 1 - sum(1 for d in week if d != 0)
                paragraph.add_run(f"{next_month_abbr}-{next_day}\n")
                shading_elm = parse_xml(r'<w:shd {} w:fill="D3D3D3"/>'.format(nsdecls('w')))
                cell._tc.get_or_add_tcPr().append(shading_elm)
            else:  # Current month's days
                paragraph.add_run(f"{month_abbr}-{day}\n")
            
            # Add two text sections for work schedule
            paragraph.add_run("____:____\n")
            paragraph.add_run("____:____")
    
    # Set font properties
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(9.5)
                    run.bold = True
    
    # Set row height (except first row)
    for row in table.rows[1:]:
        row.height = Inches(0.5)
    
    # Save document
    filename = f"BBYCAL_{month_name}_{year}.docx"
    doc.save(filename)
    return filename

def main():
    try:
        year = int(input("Enter the Year (e.g., 2025): "))
        month = int(input("Enter Month Number (1-12) which you want to generate: "))
        if 1 <= month <= 12:
            filename = create_month_calendar(year, month)
            print(f"Calendar Saved as {filename}")
        else:
            print("Please enter a valid month (1-12)")
    except ValueError:
        print("Please enter valid numbers for year and month")

if __name__ == "__main__":
    main()
